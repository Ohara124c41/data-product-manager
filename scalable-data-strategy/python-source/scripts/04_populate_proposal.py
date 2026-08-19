from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import json

src='upload/template-of-flyber-data-strategy-mvp.docx'; out='outputs/flyber/flyber-data-strategy-proposal-completed.docx'
d=Document(src); a=json.load(open('flyber_analysis.json'))['etl']; m=json.load(open('flyber_metrics.json'))

def setp(i,text,bold_prefix=None):
    p=d.paragraphs[i]; p.clear();
    if bold_prefix and text.startswith(bold_prefix):
        r=p.add_run(bold_prefix); r.bold=True; p.add_run(text[len(bold_prefix):])
    else: p.add_run(text)
    return p
def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def style_table(t):
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,c in enumerate(t.rows[0].cells):
        shade(c,'D9EAF7'); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in c.paragraphs:
            for r in p.runs: r.bold=True; r.font.color.rgb=RGBColor(23,54,93); r.font.size=Pt(8.5)
    for row in t.rows[1:]:
        for c in row.cells:
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in c.paragraphs:
                p.paragraph_format.space_after=Pt(2)
                for r in p.runs: r.font.size=Pt(8.5)
def add_image_at(i,path,width=6.4):
    p=d.paragraphs[i]; p.clear(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(); r.add_picture(path,width=Inches(width))

# Section 1
t=d.tables[0]
stakeholders=[
['Engineering / Platform','Owns the live application, instrumentation, pipeline reliability, and capacity. It is primary because every other team depends on complete, timely, trustworthy data.','1) Monitor event volume and ingestion failures to scale infrastructure. 2) Diagnose app and pipeline issues by device, page, location, and release.'],
['Product Management','Owns rider experience and prioritizes the MVP roadmap. It is primary because behavioral evidence is needed to identify friction and validate improvements.','1) Measure the rider funnel from open/search to request/begin ride. 2) Compare adoption and conversion by device, page, and neighborhood.'],
['Marketing / Growth','Owns customer acquisition and campaign investment. It is primary because campaigns directly change demand, event volume, and infrastructure load.','1) Attribute traffic and ride-intent changes to campaigns. 2) Compare campaign response by acquisition channel, location, and cohort.'],
['Finance / Operations','Owns unit economics, forecasts, and operating capacity. It is primary because transactional growth drives revenue, service demand, and warehouse cost.','1) Forecast ride demand and data-processing cost. 2) Reconcile request/begin-ride activity with operational and financial reporting.']]
for r,row in enumerate(stakeholders,1):
    for c,v in enumerate(row): t.cell(r,c).text=v
style_table(t)

t=d.tables[1]
requirements=[
['Engineering / Platform','Capacity monitoring; incident diagnosis','event_id, event_time, event_type, ingestion_time, pipeline_status; device_type, app_version, page_type, error_code','Timestamps and IDs quantify throughput and latency; status/error fields locate failures; device, release, and page dimensions isolate affected segments.'],
['Product Management','Funnel measurement; segment comparison','rider_id, session_id, event_type, event_time; page_type, device_type, neighborhood, app_version','Rider/session keys preserve journeys; ordered event types measure conversion; product dimensions reveal where behavior differs.'],
['Marketing / Growth','Campaign attribution; cohort response','rider_id, campaign_id, acquisition_channel, event_time; event_type, neighborhood, device_type, experiment_variant','Campaign/channel and time connect exposure to behavior; event and segment fields measure incremental response and support cohort comparisons.'],
['Finance / Operations','Demand and cost forecasts; transaction reconciliation','ride_id, rider_id, request_time, begin_ride_time, ride_status; event_count, bytes_processed, neighborhood, warehouse_cost','Ride keys and states support auditable transaction counts; volume, processing, location, and cost fields support capacity and expense forecasts.']]
for r,row in enumerate(requirements,1):
    for c,v in enumerate(row): t.cell(r,c).text=v
style_table(t)

schemas=[
('Rider', ['rider_id (PK)','created_at','home_neighborhood','acquisition_channel','consent_status','campaign_id'], 'rider_id is a stable surrogate identifier that avoids mutable or sensitive natural keys such as email or phone. campaign_id is nullable because organic riders have no campaign. This table stores one row per rider, preventing repeated rider attributes in every event.'),
('Session', ['session_id (PK)','rider_id (FK)','started_at','device_type','app_version','ended_at'], 'session_id uniquely identifies one app visit and supports ordered journey analysis. rider_id is the foreign key to Rider because one rider can create many sessions. Device and app-version attributes belong at session grain instead of being redundantly repeated as rider attributes.'),
('Event', ['event_id (PK)','session_id (FK)','event_time','event_type','page_type','neighborhood'], 'event_id is the immutable unique event UUID and supports deduplication. session_id links each event to its session, creating a one-to-many Session-to-Event relationship. Event time, type, page, and observed location belong at event grain; request_car and begin_ride event types support transactional funnel analysis.')]
for idx,(name,headers,rationale) in enumerate(schemas):
    setp([23,31,40][idx],name)
    tb=d.tables[2+idx]
    while len(tb.rows[0].cells)<6: tb.add_column(Inches(1.0))
    for j,h in enumerate(headers): tb.cell(0,j).text=h
    style_table(tb)
    setp([27,35,44][idx],rationale)

# Section 3 steps and ETL tables
steps=[
('1. Preserve and inspect the source. ','Used the supplied XLSX copy rather than editing the raw log in place; confirmed 124,980 event rows and the expected schema. XLSX preserves timestamp types and UUID text without CSV delimiter or encoding ambiguity.'),
('2. Standardize fields. ','Converted Excel serial event_time values to calendar dates, retained UUIDs as text, and standardized machine-readable category labels for presentation. This makes grouping reproducible while protecting identifier integrity.'),
('3. Apply the complete-day filter. ','Limited reporting to October 5-11, 2019. A partial October 12 boundary segment was excluded because comparing an incomplete day with full days would create a false decline.'),
('4. Aggregate and validate. ','Counted event_uuid by date, event type, device type, page type, and neighborhood. Each categorical subtotal was reconciled to the daily total; all seven complete days passed.')]
for i,(h,x) in zip([60,62,64,66],steps): setp(i,h+x,bold_prefix=h)
for i in [61,63,65,67]: setp(i,'')
dates=a['dates'][:7]
tables=[d.tables[5],d.tables[6],d.tables[7],d.tables[8],d.tables[9]]
for j,dt in enumerate(dates,1): tables[0].cell(0,j).text=f"{int(dt[5:7])}/{int(dt[8:10])}/{dt[:4]}"
for j,v in enumerate(a['total'][:7],1): tables[0].cell(1,j).text=f'{v:,}'
for tb,vals in [(tables[1],a['events']['values']),(tables[2],a['devices']['values']),(tables[3],a['pages']['values']),(tables[4],a['locations']['values'])]:
    for j,dt in enumerate(dates,1): tb.cell(0,j).text=f"{int(dt[5:7])}/{int(dt[8:10])}/{dt[:4]}"
    for r,row in enumerate(vals,1):
        for j,v in enumerate(row[:7],1): tb.cell(r,j).text=f'{v:,}'
for tb in tables: style_table(tb)
setp(101,'The manual ETL was useful as an MVP validation because it exposed the schema, boundary-day issue, category cardinalities, and reconciliation rules. It was not efficient or scalable: repeated full-file extraction, spreadsheet grouping, and manual copying create latency, inconsistent logic, and error risk as event volume grows. Flyber should automate an incremental pipeline that lands immutable logs in cloud object storage, validates schema and required fields, deduplicates on event_id, converts timestamps to UTC, quarantines invalid records, aggregates by partitioned event_date, and loads warehouse tables on a schedule. Orchestration, monitoring, lineage, retries, and reconciliation tests should alert Engineering before downstream dashboards refresh.')

# Section 4
setp(129,'Chosen criterion: How many events of each event type per day? This single dataset provides the most information because its row totals quantify event-log growth while the categories distinguish customer engagement (open, search, choose car) from transaction-intent and fulfilled-activity proxies (request car and begin ride). In the Section 3 sample, the category sum reconciles to total events every day, so the breakdown preserves the total and adds business meaning. Customer records themselves cannot be counted without a rider identifier or new-account event; therefore open/search/choose-car growth should be labeled a behavioral proxy, not a unique-customer count. Transactional growth is best approximated by request_car and begin_ride, while event-log growth is measured directly by the daily category sum. Event Log Data is consequently the most important source for this question, supplemented later by Customer and Ride tables for exact unique-customer and financial transaction measures.')

# Section 5 visuals
add_image_at(141,'outputs/flyber/charts/choose_car_trend.png',6.3)
setp(145,'Choose Car events rose from 61,416 to 1,875,536, a 30.54x increase. Growth accelerated into the October campaign, peaked on October 4, and stabilized at a level far above the beginning of the month. This is the fastest endpoint growth multiple among the five event types and indicates both stronger rider intent and substantial burst capacity requirements.')
setp(147,'1. Loaded the Section 5 workbook and interpreted the first column as daily dates.')
setp(148,'2. Selected Date and Choose Car, plotted a continuous line, formatted the y-axis as event count, and shaded October 1-7 as the campaign window.')
setp(149,'3. Checked the first, final, and peak values against the source table and calculated final divided by initial volume.')
add_image_at(153,'outputs/flyber/charts/begin_ride_trend.png',6.3)
setp(157,'Begin Ride events increased from 6,838 to 75,956, an 11.11x increase. The series shares the broader campaign pattern but grows more slowly than upper-funnel events, suggesting that attention and ride consideration expanded faster than completed ride starts. Product and Operations should monitor conversion from Request Car to Begin Ride as scale increases.')
setp(159,'1. Selected Date and Begin Ride from the same validated daily event-type dataset.')
setp(160,'2. Created a continuous line with consistent date range, event-count units, and campaign-window shading.')
setp(161,'3. Compared the first, final, and peak observations and interpreted the line as a lower-funnel operational measure.')

# Section 6
add_image_at(170,'outputs/flyber/charts/total_event_growth.png',6.4)
setp(174,'Daily total events increased from 790,329 on September 11 to 12,788,264 on October 11. Calculation: 12,788,264 / 790,329 = 16.18x, equivalent to an absolute increase of 11,997,935 events or 1,518.1%. Volume peaked at 18,918,096 on October 4. Because the calculation compares daily run rates exactly one month apart, it measures growth in operational load without conflating it with the number of days in a period.')
setp(178,'Event logs are the fastest-growing data class in absolute volume because every rider journey produces multiple events: open, search, choose car, request car, and begin ride. Among observed event types, Choose Car has the fastest endpoint growth at 30.54x; Open remains the largest volume at 8,438,774 events on October 11. Customer data should grow closer to unique rider acquisition, while transactional data grows with requests and rides and therefore produces fewer rows than behavioral telemetry. Exact cross-class rates require unique rider and ride records; the event dataset supports this ordering through event-type proxies, not exact customer-table counts.')
add_image_at(184,'outputs/flyber/charts/all_event_types_log.png',6.4)
setp(193,'Graph pattern: All five series broadly rise together, peak on October 4, decline through October 8, and then stabilize or increase slightly. The logarithmic scale shows that the shared shape persists despite order-of-magnitude differences in volume. Good or bad: synchronized growth is positive evidence of end-to-end engagement, but upper-funnel Choose Car growth outpaces Begin Ride growth, so conversion and supply constraints require monitoring. October campaign: the October 1-7 campaign coincides with rapid acceleration and the peak; average total volume during the campaign was 14.22 million events per day versus 4.56 million during September 24-30, a 211.6% increase. Impact: campaign activity materially increases both business demand and telemetry load, although the observational data alone does not prove causality. Importance: Marketing, Engineering, Product, and Operations must jointly forecast campaign traffic, pre-scale ingestion and app infrastructure, monitor conversion, and distinguish durable post-campaign growth from a temporary spike.')

# Section 7
setp(223,'Recommendation: use a cloud data warehouse. Cost: a startup avoids large capital purchases and can align spending with storage and query use, with budgets, retention tiers, and workload controls. Scalability: the observed 16.18x daily event growth and campaign peak require elastic compute and date-partitioned storage without lengthy hardware procurement. Expertise: managed services reduce database administration and allow the small team to focus on data models, quality, and products, although SQL, governance, security, and FinOps skills remain necessary. Latency/connectivity: a cloud warehouse supports frequent batch or near-real-time analytics close to cloud object storage; the live application should continue to use an operational database rather than querying the warehouse. Reliability: multi-zone managed infrastructure, backups, monitoring, and service-level commitments are stronger and faster to operationalize than a startup-owned cluster. An on-premise system would create fixed capacity, procurement delay, and specialized operating burden precisely when traffic is volatile.')
setp(229,'Suggested DWH: Google BigQuery. BigQuery is a serverless cloud warehouse suited to rapidly growing event data because compute scales without cluster sizing, storage and compute are separated, and standard SQL supports Engineering, Product, Marketing, Finance, and BI users. Flyber should partition fact_event by event_date, cluster by event_type and session_id or rider_id, materialize frequently used daily aggregates, and apply workload labels, query quotas, budgets, and lifecycle rules to control cost. Streaming or micro-batch ingestion can support low-latency operational dashboards, while scheduled transformations provide stable executive metrics. Built-in replication, managed maintenance, encryption, access controls, audit logs, and integration with cloud storage reduce the reliability and expertise burden. This recommendation follows directly from the correlated campaign spike, 18.9 million-event peak day, and 16.18x monthly run-rate growth; it provides elasticity while preserving a normalized core model and governed analytical marts.')

# Compact placeholder cleanup and metadata
for p in d.paragraphs:
    if '[Insert Response Here.]' in p.text or '[Insert Visualization Here.]' in p.text: p.text=''
    for r in p.runs:
        r.font.name='Open Sans'; r._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),'Open Sans'); r._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'),'Open Sans')
d.core_properties.title='Flyber Data Strategy MVP'; d.core_properties.subject='Completed data strategy proposal'; d.core_properties.author='Ohara'
Path(out).parent.mkdir(parents=True,exist_ok=True); d.save(out); print(out)
